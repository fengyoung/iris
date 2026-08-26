"""Markdown → 多格式导出转换器。"""

from __future__ import annotations

from pathlib import Path
from typing import Optional


def convert_report(markdown: str, output_path: Path, *, format: str = "md", title: str = "") -> Path:
    """将 Markdown 报告转换为指定格式。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if format == "md":
        from iris.utils.shared import atomic_write_text

        output_path = output_path.with_suffix(".md")
        atomic_write_text(output_path, markdown)
        return output_path

    if format == "docx":
        return _convert_to_docx(markdown, output_path, title=title)

    raise ValueError(f"不支持的输出格式: {format}")


def _convert_to_docx(markdown: str, output_path: Path, *, title: str = "") -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise RuntimeError("python-docx 未安装，请运行 pip install python-docx")

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "等线"
    font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped and stripped[0].isdigit() and ". " in stripped[:4]:
            doc.add_paragraph(stripped, style="List Number")
        elif stripped:
            doc.add_paragraph(stripped)
        # 空行跳过

    from iris.utils.shared import atomic_write_bytes
    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)
    atomic_write_bytes(output_path, buffer.getvalue())
    return output_path
