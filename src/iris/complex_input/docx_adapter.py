"""DOCX 文件适配器 — 提取文字内容，供纯文本流水线使用。

设计:
  - 段落文字提取：遍历 document.paragraphs
  - 表格文字提取：遍历 document.tables → 逐单元格提取
  - 嵌入图片检测：判断是否有 inline shapes（暂不提取图片，仅标记）
  - 截断：超长文档截断以避免超出 LLM 上下文窗口

用法:
    adapter = DocxAdapter()
    content = adapter.process(docx_path)
    # content.text → 提取的全部文字
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from iris.core.exceptions import IrisRuntimeError

logger = logging.getLogger(__name__)

_DEFAULT_MAX_TEXT_CHARS = 8000

# OOXML namespace 常量（用于检测嵌入图片）
_NS_DRAWING = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
_NS_PICT = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'


class DocxAdapterError(IrisRuntimeError):
    """DOCX 适配器相关错误。"""


@dataclass
class DocxContent:
    """单个 DOCX 文件的处理结果。"""

    path: str
    text: str                       # 提取的全部文字（已截断）
    paragraph_count: int = 0
    table_count: int = 0
    has_images: bool = False
    error: Optional[str] = None


class DocxAdapter:
    """DOCX 文件文字提取器。

    用法:
        adapter = DocxAdapter()
        content = adapter.process(docx_path)
    """

    def __init__(self):
        self._check_dependency()

    @staticmethod
    def _check_dependency() -> None:
        """检查 python-docx 是否可用。"""
        try:
            import docx  # noqa: F401
        except ImportError:
            raise DocxAdapterError(
                "python-docx 未安装，请运行: pip install python-docx"
            )

    # ── 公开 API ──────────────────────────────────────────────

    def process(
        self,
        docx_path: str | Path,
        *,
        max_text_chars: int = _DEFAULT_MAX_TEXT_CHARS,
    ) -> DocxContent:
        """从 DOCX 文件提取文字内容。

        Args:
            docx_path: DOCX 文件路径
            max_text_chars: 提取文字的最大字符数

        Returns:
            DocxContent 包含提取的文字和元信息
        """
        import docx as _docx

        path = Path(docx_path).resolve()
        if not path.exists():
            raise DocxAdapterError(f"DOCX 文件不存在: {path}")

        try:
            doc = _docx.Document(str(path))
        except Exception as exc:
            raise DocxAdapterError(f"无法打开 DOCX 文件: {exc}") from exc

        try:
            # ── 1. 提取段落文字 ──────────────────────────────
            text_parts: List[str] = []
            paragraph_count = 0

            for para in doc.paragraphs:
                para_text = para.text.strip()
                paragraph_count += 1
                if para_text:
                    # 按样式分级：标题加 ## 前缀
                    if para.style and para.style.name and para.style.name.startswith("Heading"):
                        try:
                            level = int(para.style.name.split()[-1])
                            prefix = "#" * min(level, 4) + " "
                        except (ValueError, IndexError):
                            prefix = ""
                        text_parts.append(f"{prefix}{para_text}")
                    else:
                        text_parts.append(para_text)

            # ── 2. 提取表格文字 ──────────────────────────────
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_lines: List[str] = [f"\n[表格 {table_count}]"]
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append(" | ".join(cells))
                    if row_idx >= 20:  # 大表格截断
                        table_lines.append("...（表格行数过多，已截断）")
                        break
                text_parts.extend(table_lines)

            # ── 3. 检测嵌入图片 ──────────────────────────────
            has_images = False
            try:
                for para in doc.paragraphs:
                    if para._element.findall(f'.//{_NS_DRAWING}'):
                        has_images = True
                        break
                    if para._element.findall(f'.//{_NS_PICT}'):
                        has_images = True
                        break
            except Exception:
                pass  # 图片检测失败不影响文字提取

            # ── 4. 拼接 + 截断 ───────────────────────────────
            full_text = "\n\n".join(text_parts)
            if len(full_text) > max_text_chars:
                full_text = full_text[:max_text_chars] + "\n\n...（文字已截断）"

            return DocxContent(
                path=str(path),
                text=full_text,
                paragraph_count=paragraph_count,
                table_count=table_count,
                has_images=has_images,
            )

        except Exception as exc:
            error_msg = f"DOCX 处理异常: {exc}"
            logger.warning(error_msg)
            return DocxContent(
                path=str(path),
                text="",
                error=error_msg,
            )
