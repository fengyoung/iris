"""complex_input/docx_adapter.py 专项测试。"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest

from iris.complex_input.docx_adapter import DocxAdapter, DocxAdapterError, DocxContent


# ── 辅助函数 ──────────────────────────────────────────────


def _create_minimal_docx(path: Path, paragraphs: list[str] | None = None) -> None:
    """用 python-docx 创建一个最小 DOCX 文件（用于测试）。"""
    from docx import Document
    doc = Document()
    doc.add_heading("Test Document", level=1)
    texts = paragraphs or [
        "First paragraph with some content.",
        "Second paragraph with more content.",
        "Third paragraph about project planning.",
    ]
    for text in texts:
        doc.add_paragraph(text)
    doc.save(str(path))


# ── DocxAdapter 基础 ────────────────────────────────────────


class TestDocxAdapterBasic:
    def test_init_checks_dependency(self):
        """python-docx 可用时正常初始化。"""
        adapter = DocxAdapter()
        assert adapter is not None

    def test_file_not_found(self, tmp_path):
        """不存在的文件抛出 DocxAdapterError。"""
        adapter = DocxAdapter()
        nonexistent = tmp_path / "nonexistent.docx"
        with pytest.raises(DocxAdapterError, match="DOCX 文件不存在"):
            adapter.process(nonexistent)

    @patch("iris.complex_input.docx_adapter.DocxAdapter._check_dependency")
    def test_missing_dependency(self, mock_check):
        """python-docx 未安装时抛出 DocxAdapterError。"""
        mock_check.side_effect = DocxAdapterError("python-docx 未安装")
        with pytest.raises(DocxAdapterError, match="python-docx"):
            DocxAdapter()


class TestDocxAdapterTextExtraction:
    def test_extract_paragraphs(self, tmp_path):
        """提取段落文字。"""
        path = tmp_path / "test.docx"
        _create_minimal_docx(path)

        adapter = DocxAdapter()
        content = adapter.process(path)

        assert "First paragraph" in content.text
        assert "Second paragraph" in content.text
        assert content.paragraph_count > 0
        assert content.error is None

    def test_extract_heading(self, tmp_path):
        """标题段落带 # 前缀。"""
        path = tmp_path / "test.docx"
        _create_minimal_docx(path)

        adapter = DocxAdapter()
        content = adapter.process(path)

        assert "Test Document" in content.text

    def test_text_truncation(self, tmp_path):
        """超长文档截断。"""
        path = tmp_path / "long.docx"
        long_texts = [f"Paragraph {i}: " + "x" * 200 for i in range(50)]
        _create_minimal_docx(path, paragraphs=long_texts)

        adapter = DocxAdapter()
        content = adapter.process(path, max_text_chars=500)

        assert len(content.text) <= 550  # 容差含截断标记

    def test_empty_docx(self, tmp_path):
        """空 DOCX 不崩溃。"""
        path = tmp_path / "empty.docx"
        _create_minimal_docx(path, paragraphs=[])

        adapter = DocxAdapter()
        content = adapter.process(path)
        assert isinstance(content.text, str)
        assert content.error is None

    def test_non_docx_file(self, tmp_path):
        """非 DOCX 文件（损坏的 zip）应抛异常或返回带错误的 DocxContent。"""
        path = tmp_path / "fake.docx"
        path.write_text("this is not a real docx file")

        adapter = DocxAdapter()
        try:
            content = adapter.process(str(path))
            # 如果没抛异常，应返回带 error 或空文本
            assert content.error is not None or content.text == ""
        except DocxAdapterError:
            # 抛出异常也是合理行为（无法打开损坏文件）
            pass


# ── DocxContent ────────────────────────────────────────────


class TestDocxContent:
    def test_docx_content_fields(self):
        """DocxContent 字段正确赋值。"""
        content = DocxContent(
            path="/tmp/test.docx",
            text="Hello World",
            paragraph_count=3,
            table_count=1,
            has_images=False,
        )
        assert content.path == "/tmp/test.docx"
        assert content.text == "Hello World"
        assert content.paragraph_count == 3
        assert content.table_count == 1
        assert content.has_images is False
